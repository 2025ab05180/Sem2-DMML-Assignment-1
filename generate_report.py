"""
RecoMart - PDF Report Generator (Word Document)
Generates the complete submission report as a .docx file with all required sections.
"""

import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(__file__).resolve().parent
SCREENSHOTS_DIR = BASE_DIR / "submission_package" / "screenshots"
REPORTS_DIR = BASE_DIR / "reports"


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return h


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def add_screenshot(doc, filename, caption="", width=6.0):
    img_path = SCREENSHOTS_DIR / filename
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        doc.add_paragraph(f"[Screenshot not found: {filename}]")


def load_latest_report(pattern):
    files = sorted(REPORTS_DIR.glob(pattern), reverse=True)
    if files:
        with open(files[0]) as f:
            return json.load(f)
    return {}


def generate_report():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ========== TITLE PAGE ==========
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("End-to-End Data Management Pipeline\nfor a Recommendation System")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("RecoMart — E-Commerce Recommendation Engine")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_lines = [
        "Data Management for Machine Learning (AIMLCZG529 / DSECLZG529)",
        "Semester: S2-25",
        "Assignment I — EC1 (20 Marks)",
        "Group Number: 43",
        "Programme: M.Tech AI/ML — BITS Pilani WILP",
        "Submission Date: 22.07.2026",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # ========== TEAM MEMBER DETAILS ==========
    team_heading = doc.add_paragraph()
    team_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = team_heading.add_run("Team Member Details")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    doc.add_paragraph("")

    add_table(doc,
              ["Name", "BITS ID", "Email"],
              [
                  ["BRAJESH MISHRA", "2025AB05161", "2025ab05161@wilp.bits-pilani.ac.in"],
                  ["YENGANTIWAR PRASHANT SAINATH", "2025AB05178", "2025ab05178@wilp.bits-pilani.ac.in"],
                  ["ARTHIKA G", "2025AB05180", "2025ab05180@wilp.bits-pilani.ac.in"],
                  ["A SRIKAR", "2025AB05185", "2025ab05185@wilp.bits-pilani.ac.in"],
                  ["GIRIDHARAN B", "2025AB05188", "2025ab05188@wilp.bits-pilani.ac.in"],
              ])

    doc.add_page_break()

    # ========== TABLE OF CONTENTS ==========
    add_heading_styled(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Problem Statement",
        "2. Objectives",
        "3. Data Sources",
        "4. Methodology / Pipeline Architecture",
        "5. Implementation Details",
        "   5.1 Data Generation & Ingestion",
        "   5.2 Raw Data Storage",
        "   5.3 Data Validation & Quality Report",
        "   5.4 Data Preparation & EDA",
        "   5.5 Feature Engineering & Transformation",
        "   5.6 Feature Store",
        "   5.7 Data Versioning & Lineage",
        "   5.8 Model Training & Evaluation",
        "   5.9 Pipeline Orchestration",
        "6. Results and Output Screenshots",
        "7. Conclusion and Future Scope",
        "8. Google Drive Links",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ========== 1. PROBLEM STATEMENT ==========
    add_heading_styled(doc, "1. Problem Statement", level=1)
    doc.add_paragraph(
        "RecoMart is an e-commerce startup that wants to build a data-driven product recommendation engine "
        "to improve customer engagement and cross-selling opportunities. The platform collects user behaviour "
        "data from multiple channels — web/mobile clickstream logs, transactional purchase history, product "
        "catalog metadata, and external API data."
    )
    doc.add_paragraph(
        "The goal of this project is to design and implement a complete, automated, modular data management "
        "pipeline that continuously ingests raw data, validates and cleans it, engineers features, trains "
        "recommendation models, and tracks every artefact along the way — following the practices of modern "
        "ML data stacks."
    )
    doc.add_paragraph(
        "The business objective is to increase conversion rates by providing personalized product "
        "recommendations to users based on their browsing behaviour, purchase history, and explicit ratings."
    )

    # ========== 2. OBJECTIVES ==========
    add_heading_styled(doc, "2. Objectives", level=1)
    objectives = [
        "Ingest data from heterogeneous sources (CSV files + REST APIs) with error handling, retry logic, and audit logging.",
        "Store raw data in a structured, partitioned data lake layout.",
        "Profile and validate data quality (missing values, duplicates, schema mismatches, range violations).",
        "Clean and preprocess data — handle missing values, encode categoricals, normalise numerics.",
        "Engineer user-level, item-level, and user–item interaction features suitable for collaborative and content-based filtering.",
        "Manage features through a versioned feature store with a metadata registry.",
        "Version all datasets (raw, processed, features) with SHA-256 integrity tracking and a manifest.",
        "Train and evaluate recommendation models (SVD collaborative filtering + content-based cosine similarity).",
        "Track experiments, parameters, and metrics with MLflow.",
        "Orchestrate the full pipeline as a DAG using Prefect (with a sequential fallback).",
    ]
    for i, obj in enumerate(objectives, 1):
        doc.add_paragraph(f"{i}. {obj}")

    # ========== 3. DATA SOURCES ==========
    add_heading_styled(doc, "3. Data Sources", level=1)
    doc.add_paragraph(
        "The pipeline ingests data from multiple heterogeneous sources, simulating a real-world e-commerce "
        "environment. Synthetic data is generated with intentional quality issues (missing values, duplicates, "
        "out-of-range ratings) to demonstrate the validation and cleaning pipeline."
    )
    add_table(doc,
              ["Source", "Type", "Records", "Description"],
              [
                  ["Users", "CSV", "1,000", "Demographics — age, gender, country, signup date, premium status"],
                  ["Products", "CSV + JSON (API)", "500", "Catalog — name, category, price, brand, avg rating"],
                  ["Clickstream", "CSV", "~51,000", "Events — views, clicks, add-to-cart, wishlist, search"],
                  ["Transactions", "CSV", "10,000", "Purchases — quantities, amounts, payment methods"],
                  ["Ratings", "CSV", "~20,000", "Explicit user–item ratings (1–5 scale) with review text"],
                  ["External API", "REST (FakeStore)", "20", "Supplementary product data from live API"],
              ])

    doc.add_page_break()

    # ========== 4. METHODOLOGY / PIPELINE ==========
    add_heading_styled(doc, "4. Methodology / Pipeline Architecture", level=1)
    doc.add_paragraph(
        "The pipeline follows a modular, stage-based architecture where each component operates independently "
        "and produces structured outputs consumed by the next stage. The full DAG is:"
    )
    doc.add_paragraph(
        "Data Generation → Data Ingestion → Data Validation → Data Preparation → "
        "Feature Engineering → Feature Store → Data Versioning → Model Training",
        style="Intense Quote"
    )
    doc.add_paragraph(
        "All stages are orchestrated by Prefect (with a sequential fallback) and produce timestamped "
        "logs and JSON reports for auditability."
    )
    add_screenshot(doc, "12_folder_structure.png", "Figure 1: Data Lake Storage Structure", width=5.0)

    doc.add_page_break()

    # ========== 5. IMPLEMENTATION DETAILS ==========
    add_heading_styled(doc, "5. Implementation Details", level=1)

    # 5.1
    add_heading_styled(doc, "5.1 Data Generation & Ingestion", level=2)
    doc.add_paragraph(
        "Stage 1 generates realistic synthetic e-commerce data using the Faker library with fixed seeds "
        "for reproducibility. It intentionally injects ~5% missing values in user ages, ~3% in genders, "
        "~2% missing transaction amounts, ~2% duplicate clickstream events, and ~1% out-of-range ratings."
    )
    doc.add_paragraph(
        "Stage 2 ingests data from two source types: CSV files read from the local filesystem, and product "
        "data fetched from the FakeStore REST API (https://fakestoreapi.com/products). The API ingestion "
        "uses exponential-backoff retry (up to 3 attempts) with a 30-second timeout. All ingestion events "
        "are logged with Loguru and a JSON ingestion report is saved per run."
    )
    doc.add_paragraph("Key files: src/ingestion/generate_data.py, src/ingestion/ingest.py")

    # 5.2
    add_heading_styled(doc, "5.2 Raw Data Storage", level=2)
    doc.add_paragraph(
        "Ingested data is stored in a local data lake under data/raw/, structured as:"
    )
    doc.add_paragraph(
        "data/raw/{source_name}/{source}_{YYYYMMDD}.csv — partitioned by source and date-stamped filenames. "
        "This mirrors production data lake patterns (e.g., S3 prefix-based partitioning)."
    )
    doc.add_paragraph("Storage structure is documented in the project README and pipeline_config.yaml.")

    # 5.3
    add_heading_styled(doc, "5.3 Data Validation & Quality Report", level=2)
    doc.add_paragraph(
        "The validation module performs four categories of checks on every dataset:"
    )
    checks = [
        "Schema validation: Verifies required columns exist and detects extra/missing columns.",
        "Missing value analysis: Per-column and overall missing percentages.",
        "Duplicate detection: Row-level duplicate counts and percentages.",
        "Range validation: Rating values must be in [1–5], prices non-negative, quantities positive.",
    ]
    for c in checks:
        doc.add_paragraph(c, style="List Bullet")
    doc.add_paragraph(
        "A comprehensive JSON quality report is generated per run and saved to reports/. "
        "Validated copies of each dataset are saved to data/validated/."
    )
    doc.add_paragraph("Key file: src/validation/validate.py")
    add_screenshot(doc, "02_data_quality_report.png", "Figure 2: Data Quality Validation Report", width=6.0)

    doc.add_page_break()

    # 5.4
    add_heading_styled(doc, "5.4 Data Preparation & Exploratory Data Analysis", level=2)
    doc.add_paragraph("The preparation module performs the following cleaning and preprocessing steps:")
    prep_steps = [
        "Missing value imputation: Median for ages, recalculated for transaction amounts, 'Unknown' for categoricals.",
        "Categorical encoding: LabelEncoder for gender, product category, event type, device, payment method.",
        "Numerical normalisation: MinMaxScaler for product prices.",
        "Derived features: Account age in days, hour-of-day and day-of-week from timestamps.",
        "Filtering: Removes out-of-range ratings (< 1 or > 5), deduplicates clickstream events.",
    ]
    for s in prep_steps:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Key file: src/preparation/prepare.py, notebooks/eda_analysis.ipynb")

    add_screenshot(doc, "03_eda_users.png", "Figure 3: User Demographics — Age, Gender, Country Distribution", width=6.0)
    add_screenshot(doc, "04_eda_products.png", "Figure 4: Product Catalog — Price, Category, Rating Distribution", width=6.0)

    doc.add_page_break()

    add_screenshot(doc, "05_eda_clickstream.png", "Figure 5: Clickstream — Event Types, Devices, Hourly Activity", width=6.0)
    add_screenshot(doc, "08_eda_transactions.png", "Figure 6: Transactions — Status, Payment Methods, Amounts", width=6.0)
    add_screenshot(doc, "06_eda_ratings.png", "Figure 7: Ratings — Distribution, Per-User, Per-Item Counts", width=6.0)

    doc.add_page_break()

    add_screenshot(doc, "07_interaction_heatmap.png", "Figure 8: User-Item Interaction Heatmap (Top 25 × 25)", width=5.5)

    # Load sparsity info
    eda_data = load_latest_report("eda_summary_*.json")
    if eda_data and "ratings" in eda_data:
        sparsity = eda_data["ratings"].get("sparsity", 0)
        doc.add_paragraph(
            f"The user-item interaction matrix has a sparsity of {sparsity:.2%}, which is typical for "
            "recommendation system datasets. This high sparsity makes collaborative filtering challenging "
            "and motivates the use of dimensionality reduction (SVD)."
        )

    doc.add_page_break()

    # 5.5
    add_heading_styled(doc, "5.5 Feature Engineering & Transformation", level=2)
    doc.add_paragraph("Three feature tables are created from the cleaned data:")

    doc.add_paragraph("User Features (15 features):", style="List Bullet")
    doc.add_paragraph(
        "total_ratings, avg_rating, rating_std, total_purchases, total_spend, avg_order_value, "
        "total_clicks, view_count, cart_add_count, purchase_frequency, activity_score, "
        "account_age_days, is_premium, gender_encoded"
    )
    doc.add_paragraph("Item Features (13 features):", style="List Bullet")
    doc.add_paragraph(
        "total_ratings, avg_user_rating, rating_variance, total_purchases, total_revenue, "
        "total_views, view_to_purchase_ratio, popularity_score, price, price_normalized, "
        "category_encoded, in_stock"
    )
    doc.add_paragraph("User-Item Features (7 features):", style="List Bullet")
    doc.add_paragraph(
        "rating, n_interactions, has_viewed, has_carted, has_purchased"
    )
    doc.add_paragraph(
        "A SQL schema (data/features/schema.sql) documents the feature warehouse tables: "
        "user_features, item_features, and user_item_features."
    )
    doc.add_paragraph("Key file: src/transformation/transform.py")

    # 5.6
    add_heading_styled(doc, "5.6 Feature Store", level=2)
    doc.add_paragraph(
        "A custom feature store is implemented using SQLite as a metadata registry. It provides:"
    )
    store_features = [
        "Feature set registration with versioned snapshots (data/feature_store/versions/).",
        "Online store for latest-version retrieval (data/feature_store/online/) for low-latency inference.",
        "Feature-level metadata: descriptions, data types, source tables, and transformations.",
        "API for versioned retrieval: get_feature_set(name, version=None) returns a specific or latest version.",
    ]
    for s in store_features:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Key file: src/feature_store/store.py")

    add_screenshot(doc, "10_feature_store.png", "Figure 9: Feature Store Registry & Metadata", width=6.0)

    doc.add_page_break()

    # 5.7
    add_heading_styled(doc, "5.7 Data Versioning & Lineage", level=2)
    doc.add_paragraph(
        "All datasets (raw, processed, features) are versioned using SHA-256 file hashing. "
        "A manifest.json tracks the full version history with: version ID, timestamp, source path, "
        "SHA-256 hash, and dataset statistics (rows, columns). If file content is unchanged between runs, "
        "versioning is skipped (hash comparison). 13 datasets are versioned across three tiers."
    )
    doc.add_paragraph("Key file: src/versioning/version.py")
    add_screenshot(doc, "11_data_versioning.png", "Figure 10: Data Versioning Manifest — 13 Datasets Tracked", width=6.0)

    doc.add_page_break()

    # 5.8
    add_heading_styled(doc, "5.8 Model Training & Evaluation", level=2)
    doc.add_paragraph("Two recommendation models are trained and evaluated:")

    doc.add_paragraph("1. SVD Collaborative Filtering", style="List Bullet")
    doc.add_paragraph(
        "Truncated SVD on the centered user–item rating matrix with 50 latent factors. "
        "The matrix is mean-centered per user before decomposition. Predicted ratings are "
        "clipped to [1, 5]."
    )
    doc.add_paragraph("2. Content-Based Filtering", style="List Bullet")
    doc.add_paragraph(
        "Cosine similarity computed on normalised item feature vectors (12 numeric features). "
        "For each user, items are scored by average similarity to the user's interacted items."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Evaluation Metrics:")

    perf = load_latest_report("model_performance_*.json")
    svd_m = perf.get("svd", {}).get("metrics", {})
    cb_m = perf.get("content_based", {}).get("metrics", {})

    add_table(doc,
              ["Metric", "SVD", "Content-Based"],
              [
                  ["Precision@5", f"{svd_m.get('precision@5', 0):.4f}", f"{cb_m.get('precision@5', 0):.4f}"],
                  ["Recall@5", f"{svd_m.get('recall@5', 0):.4f}", f"{cb_m.get('recall@5', 0):.4f}"],
                  ["NDCG@5", f"{svd_m.get('ndcg@5', 0):.4f}", f"{cb_m.get('ndcg@5', 0):.4f}"],
                  ["Precision@10", f"{svd_m.get('precision@10', 0):.4f}", f"{cb_m.get('precision@10', 0):.4f}"],
                  ["Recall@10", f"{svd_m.get('recall@10', 0):.4f}", f"{cb_m.get('recall@10', 0):.4f}"],
                  ["NDCG@10", f"{svd_m.get('ndcg@10', 0):.4f}", f"{cb_m.get('ndcg@10', 0):.4f}"],
                  ["Precision@20", f"{svd_m.get('precision@20', 0):.4f}", f"{cb_m.get('precision@20', 0):.4f}"],
                  ["Recall@20", f"{svd_m.get('recall@20', 0):.4f}", f"{cb_m.get('recall@20', 0):.4f}"],
                  ["NDCG@20", f"{svd_m.get('ndcg@20', 0):.4f}", f"{cb_m.get('ndcg@20', 0):.4f}"],
                  ["RMSE", f"{svd_m.get('rmse', 0):.4f}", "—"],
                  ["MAE", f"{svd_m.get('mae', 0):.4f}", "—"],
              ])

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Metrics are modest because the data is purely synthetic (random). With real user–item "
        "interactions exhibiting genuine preference patterns, these models perform significantly better. "
        "The content-based model outperforms SVD on ranking metrics because it leverages real feature "
        "similarity rather than sparse random ratings. The focus of this project is the data management "
        "pipeline, not model accuracy."
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "All experiments are tracked with MLflow (SQLite backend). Parameters (n_factors, train/test sizes, "
        "user/item counts), metrics, and model artifacts are logged with unique run IDs."
    )

    run_id = perf.get("mlflow_run_id", "N/A")
    doc.add_paragraph(f"MLflow Run ID: {run_id}")
    doc.add_paragraph("Key file: src/training/train.py")

    add_screenshot(doc, "09_model_performance.png", "Figure 11: Model Performance — SVD vs Content-Based", width=6.0)

    doc.add_page_break()

    add_screenshot(doc, "13_mlflow_tracking.png", "Figure 12: MLflow Experiment Tracking — Run Details", width=5.5)

    # Sample recommendations
    recs = perf.get("sample_recommendations", {})
    if recs:
        doc.add_paragraph("")
        doc.add_paragraph(f"Sample Top-5 Recommendations for User {recs.get('user_id', '?')}:")
        add_table(doc,
                  ["Rank", "Product ID", "Predicted Rating"],
                  [[i + 1, r["product_id"], r["predicted_rating"]]
                   for i, r in enumerate(recs.get("recommendations", []))])

    add_screenshot(doc, "14_sample_recommendations.png",
                   "Figure 13: Sample Recommendations Output", width=5.0)

    doc.add_page_break()

    # 5.9
    add_heading_styled(doc, "5.9 Pipeline Orchestration", level=2)
    doc.add_paragraph(
        "The entire pipeline is orchestrated using Prefect. Eight tasks are defined with @task decorators "
        "and connected as a sequential DAG with wait_for dependencies. Ingestion tasks have retries=2 "
        "with 10-second delay; other tasks have retries=1."
    )
    doc.add_paragraph(
        "When Prefect is not available, a sequential fallback runner executes all stages with "
        "try/except error handling, halting on failure. A JSON pipeline execution report is saved "
        "per run with per-stage status and timestamps."
    )
    doc.add_paragraph("Key file: src/orchestration/pipeline.py")
    doc.add_paragraph("")
    doc.add_paragraph("Pipeline DAG:")
    doc.add_paragraph(
        "generate_data → ingest_data → validate_data → prepare_data → "
        "engineer_features → feature_store → version_data → train_model",
        style="Intense Quote"
    )

    add_screenshot(doc, "01_pipeline_execution.png",
                   "Figure 14: Pipeline Execution Summary — All 8 Stages SUCCESS", width=6.0)

    doc.add_page_break()

    # ========== 6. RESULTS & SCREENSHOTS ==========
    add_heading_styled(doc, "6. Results and Output Screenshots", level=1)
    doc.add_paragraph(
        "All pipeline stages executed successfully. The following table summarises the key outputs:"
    )
    add_table(doc,
              ["Output", "Location", "Description"],
              [
                  ["Raw Data", "data/raw/", "5 sources, date-partitioned CSVs + API JSON"],
                  ["Validated Data", "data/validated/", "Post-validation copies with quality flags"],
                  ["Cleaned Data", "data/processed/", "5 clean CSVs ready for feature engineering"],
                  ["Feature Tables", "data/features/", "3 feature CSVs + SQL schema"],
                  ["Feature Store", "data/feature_store/", "SQLite registry + versioned snapshots"],
                  ["Versioned Data", "data/versions/", "13 datasets with SHA-256 manifest"],
                  ["Trained Models", "models/", "SVD + content-based .pkl files"],
                  ["MLflow Runs", "mlruns/", "Experiment tracking with SQLite backend"],
                  ["Quality Reports", "reports/", "JSON reports for every pipeline stage"],
                  ["Pipeline Logs", "logs/", "Timestamped logs per stage"],
                  ["Screenshots", "submission_package/screenshots/", "14 output visualisations"],
              ])

    doc.add_page_break()

    # ========== 7. TECHNOLOGY STACK ==========
    add_heading_styled(doc, "6.1 Technology Stack", level=2)
    add_table(doc,
              ["Component", "Tool / Library"],
              [
                  ["Language", "Python 3.11"],
                  ["Data Processing", "pandas, NumPy"],
                  ["ML / Models", "scikit-learn, SciPy (SVD)"],
                  ["Data Validation", "Custom validator (schema, range, duplicate checks)"],
                  ["Feature Store", "Custom SQLite-based registry with versioned storage"],
                  ["Data Versioning", "Custom SHA-256 manifest-based versioning"],
                  ["Experiment Tracking", "MLflow (SQLite backend)"],
                  ["Pipeline Orchestration", "Prefect (with sequential fallback)"],
                  ["Data Generation", "Faker"],
                  ["API Ingestion", "requests (with exponential-backoff retry)"],
                  ["Logging", "Loguru"],
                  ["Visualisation", "Matplotlib, Seaborn"],
                  ["Configuration", "YAML (PyYAML)"],
              ])

    doc.add_page_break()

    # ========== 7. CONCLUSION ==========
    add_heading_styled(doc, "7. Conclusion and Future Scope", level=1)

    add_heading_styled(doc, "Conclusion", level=2)
    doc.add_paragraph(
        "This project demonstrates a fully functional, modular, end-to-end data management pipeline "
        "that takes raw multi-source e-commerce data through ingestion, validation, cleaning, feature "
        "engineering, storage, versioning, model training, and evaluation — all orchestrated as a "
        "reproducible workflow. Every stage produces auditable logs and structured reports."
    )
    doc.add_paragraph(
        "The pipeline successfully processes 1,000 users, 500 products, 51,000 clickstream events, "
        "10,000 transactions, and 20,000 ratings. Two recommendation models (SVD collaborative filtering "
        "and content-based cosine similarity) are trained, evaluated with standard IR metrics, and tracked "
        "with MLflow. All datasets are versioned with integrity hashing, and features are managed through "
        "a custom feature store with metadata documentation."
    )

    add_heading_styled(doc, "Future Scope", level=2)
    future = [
        "Real data integration: Connect to actual e-commerce databases and streaming sources (Kafka).",
        "Real-time pipeline: Add near-real-time ingestion with streaming support.",
        "Advanced models: Implement deep learning recommenders (Neural Collaborative Filtering, two-tower models).",
        "A/B testing framework: Online evaluation of recommendation quality.",
        "Cloud deployment: Migrate to AWS S3 + Glue + SageMaker or GCP equivalents.",
        "Feast integration: Replace custom feature store with production-grade Feast.",
        "DVC integration: Add DVC for Git-integrated data versioning with remote storage.",
        "CI/CD: Automate pipeline testing and deployment with GitHub Actions.",
    ]
    for f in future:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_page_break()

    # ========== 8. GOOGLE DRIVE LINKS ==========
    add_heading_styled(doc, "8. Google Drive Links", level=1)

    doc.add_paragraph("")
    p1 = doc.add_paragraph()
    run = p1.add_run("Video Walkthrough (5–10 mins):")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph("<INSERT GOOGLE DRIVE LINK TO VIDEO HERE>")

    doc.add_paragraph("")
    p2 = doc.add_paragraph()
    run = p2.add_run("Project Deliverables (.zip):")
    run.bold = True
    run.font.size = Pt(12)
    doc.add_paragraph("<INSERT GOOGLE DRIVE LINK TO ZIP HERE>")

    doc.add_paragraph("")
    doc.add_paragraph(
        "Note: Both Google Drive links should be accessible to any BITS ID "
        "(sharing set to 'Anyone with BITS email can view')."
    )

    # ========== SAVE ==========
    output_path = BASE_DIR / "RecoMart_Report.docx"
    doc.save(str(output_path))
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_report()
